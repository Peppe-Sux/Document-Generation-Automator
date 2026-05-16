import os
import re
import subprocess
import shutil
import pandas as pd
from docx import Document

"""
DOCUMENT GENERATION AUTOMATOR - PROFESSIONAL VERSION
-----------------------------------------------------
Demonstrated Skills:
1. Universal Placeholder Engine: Detects and replaces custom tags ([TAG], {TAG}, <TAG>).
2. Advanced Paragraph Manipulation: Safe replacement at the paragraph text level to avoid MS Word's native XML run fragmentation bug.
3. Data Context Mapping: Automatic formatting and parsing of Pandas DataFrame records.
4. Cross-Platform PDF Export: Headless conversion via LibreOffice CLI, removing local desktop application dependencies.
5. Dynamic Naming Layer: Custom output filename mapping based on extracted row metadata.
"""

def format_context(row):
    """Sanitizes and formats Excel row data for the document template."""
    context = {}
    context['first_name'] = str(row['first_name']).strip().title()
    context['last_name'] = str(row['last_name']).strip().title()
    
    try:
        context['date'] = pd.to_datetime(row['date']).strftime('%m/%d/%Y')
    except:
        context['date'] = str(row['date'])
        
    if 'amount' in row and pd.notna(row['amount']):
        # Standard international currency formatting: 1,250.00 $ / €
        context['amount'] = f"{row['amount']:,.2f}" + " €"
    else:
        context['amount'] = "0.00 €"
        
    return context

def replace_placeholders_safely(doc_path, context, output_path, style_type="brackets"):
    """
    Replaces placeholders inside the Word document by targeting full paragraph text blocks.
    This architecture bypasses the run-splitting XML behavior native to Microsoft Word.
    """
    delimiters = {
        "brackets": (r"\[", r"\]"),          # [TAG]
        "braces": (r"\{", r"\}"),            # {TAG}
        "double_braces": (r"\{\{", r"\}\}"), # {{TAG}}
        "angles": (r"<", r">"),              # <TAG>
        "parens": (r"\(", r"\)")             # (TAG)
    }
    
    start_delim, end_delim = delimiters.get(style_type, delimiters["brackets"])
    doc = Document(doc_path)
    
    def apply_regex_to_paragraph(paragraph):
        p_text = paragraph.text
        modified = False
        for key, value in context.items():
            pattern = f"{start_delim}{key}{end_delim}"
            if re.search(pattern, p_text, flags=re.IGNORECASE):
                p_text = re.sub(pattern, value, p_text, flags=re.IGNORECASE)
                modified = True
        if modified:
            paragraph.text = p_text

    # Scan body paragraphs
    for paragraph in doc.paragraphs:
        apply_regex_to_paragraph(paragraph)
        
    # Scan tables (crucial for commercial contracts and invoices)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    apply_regex_to_paragraph(paragraph)
            
    doc.save(output_path)

def convert_docx_to_pdf(docx_path, output_dir):
    """Converts a DOCX file to PDF using LibreOffice Headless CLI (Server-Ready)."""
    paths = [
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        "/usr/bin/soffice",
        "/Applications/LibreOffice.app/Contents/MacOS/soffice"
    ]
    soffice_path = next((p for p in paths if os.path.exists(p)), None)
    
    if soffice_path:
        try:
            subprocess.run([
                soffice_path, '--headless', '--convert-to', 'pdf',
                '--outdir', output_dir, docx_path
            ], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, check=True)
            return True
        except:
            return False
    return False

def process_batch(template_file, excel_file, output_file=None, force_pdf=False, tag_style="brackets", output_dir="Output"):
    """Main batch processing pipeline for data ingestion, rendering, and routing."""
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        
    df = pd.read_excel(excel_file)
    print(f"[PROCESSING] Starting batch for {len(df)} records. Tag style: '{tag_style}'")
    
    for index, row in df.iterrows():
        context = format_context(row)
        
        # Dynamic output filename handling
        if output_file is None:
            ext = '.pdf' if force_pdf else '.docx'
            filename = f"Document_{context['first_name']}_{context['last_name']}{ext}"
        else:
            filename = output_file
            filename = filename.replace("[first_name]", context['first_name']).replace("[last_name]", context['last_name'])
            filename = filename.replace("{first_name}", context['first_name']).replace("{last_name}", context['last_name'])
            filename = filename.replace("<first_name>", context['first_name']).replace("<last_name>", context['last_name'])
            
            if force_pdf and not filename.endswith('.pdf'):
                filename = os.path.splitext(filename)[0] + '.pdf'
                
        output_path = os.path.join(output_dir, filename)
        base_output_name = os.path.splitext(output_path)[0]
        temp_docx_path = f"{base_output_name}_temp.docx"
        
        # 1. Safe compilation layer
        replace_placeholders_safely(template_file, context, temp_docx_path, style_type=tag_style)
        
        # 2. Output formatting and polymorphism routing
        if force_pdf or output_path.endswith('.pdf'):
            final_pdf_path = f"{base_output_name}.pdf"
            success = convert_docx_to_pdf(temp_docx_path, output_dir)
            if success:
                generated_pdf = temp_docx_path.replace('_temp.docx', '_temp.pdf')
                if os.path.exists(generated_pdf):
                    shutil.move(generated_pdf, final_pdf_path)
                os.remove(temp_docx_path)
                print(f" -> Generated PDF: {final_pdf_path}")
            else:
                shutil.move(temp_docx_path, f"{base_output_name}.docx")
                print(f" -> [FALLBACK] LibreOffice not found. Saved Word document: {base_output_name}.docx")
        else:
            shutil.move(temp_docx_path, f"{base_output_name}.docx")
            print(f" -> Generated Word document: {base_output_name}.docx")


# --- DYNAMIC DEMO USE CASE EXECUTION LAYER ---
if __name__ == "__main__":
    print("=== STARTING DOCUMENT AUTOMATION PIPELINE (DIRECT EXECUTION) ===\n")

    # Locate script target folder dynamically
    SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
    
    # Path routing declarations - Aligned with setup defaults
    template_path = os.path.join(SCRIPT_DIR, 'template_brackets_demo.docx')
    excel_path = os.path.join(SCRIPT_DIR, 'data_source.xlsx')
    output_folder = os.path.join(SCRIPT_DIR, 'Output_Documents')

    # Execution isolation wrapper
    def safe_run_case(case_number, description, func, *args, **kwargs):
        print(f"\n--- TEST CASE {case_number}: {description} ---")
        try:
            func(*args, **kwargs)
        except Exception as e:
            print(f"[CASE {case_number} ERROR] Execution halted: {e}")

    # =========================================================================
    # 1. Standard Batch Generation -> Word Docs (.docx)
    # =========================================================================
    safe_run_case(1, "Standard DOCX Generation using Default [brackets] Tags", 
                  process_batch, template_path, excel_path, 
                  tag_style="brackets", output_dir=output_folder)

    # =========================================================================
    # 2. Batch Generation with Conversion -> PDFs (.pdf)
    # =========================================================================
    # Using 'angles' for case 2 to match setup capabilities if needed, or keeping brackets
    safe_run_case(2, "Automated Headless PDF Compilation (Requires LibreOffice)", 
                  process_batch, template_path, excel_path, 
                  tag_style="brackets", force_pdf=True, output_dir=output_folder)

    # =========================================================================
    # 3. Custom File Naming Architecture & Alternative Tags ({braces})
    # =========================================================================
    template_braces = os.path.join(SCRIPT_DIR, 'template_braces_demo.docx')
    safe_run_case(3, "Custom Filenames using Metadata Mapping and Curly Braces Template", 
                  process_batch, template_braces, excel_path, 
                  output_file="Invoice_{last_name}_{first_name}.docx", 
                  tag_style="braces", output_dir=output_folder)
